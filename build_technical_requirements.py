from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT
from pathlib import Path

OUT = Path('/Users/wangpeng/Downloads/0626 jili/交付文件/吉利试制样车改制数字化项目招标技术要求书_V15.docx')

BLUE = '1F4E79'
DARK = '1F1F1F'
MUTED = '666666'
LIGHT = 'D9EAF7'
PALE = 'EEF5FA'
GRID = 'AAB7C4'
RED = '9C0006'

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn('w:' + m))
        if node is None:
            node = OxmlElement('w:' + m)
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)

def set_table_borders(table, color=GRID, size='4'):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = 'start' if edge == 'left' else 'end' if edge == 'right' else edge
        el = OxmlElement('w:' + tag)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), size)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)

def set_repeat_no_split(row):
    trPr = row._tr.get_or_add_trPr()
    cant = OxmlElement('w:cantSplit')
    trPr.append(cant)

def set_run_font(run, name='Microsoft YaHei', size=10.5, bold=None, color=DARK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = instruction
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.extend([fldChar1, instrText, fldChar2])

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.2)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.2)
sec.header_distance = Cm(1.0)
sec.footer_distance = Cm(1.0)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Microsoft YaHei'; normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for name, size, before, after, color in [
    ('Title', 26, 0, 8, DARK), ('Subtitle', 14, 0, 8, MUTED),
    ('Heading 1', 16, 16, 8, BLUE), ('Heading 2', 13, 12, 6, BLUE), ('Heading 3', 11.5, 8, 4, '274C77')]:
    st = styles[name]
    st.font.name = 'Microsoft YaHei'; st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    st.font.size = Pt(size); st.font.bold = name != 'Subtitle'; st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

if 'Requirement' not in styles:
    st = styles.add_style('Requirement', WD_STYLE_TYPE.PARAGRAPH)
else:
    st = styles['Requirement']
st.base_style = styles['Normal']
st.font.name = 'Microsoft YaHei'; st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
st.font.size = Pt(10.5)
st.paragraph_format.left_indent = Cm(0)
st.paragraph_format.first_line_indent = Pt(21)
st.paragraph_format.space_after = Pt(4)
st.paragraph_format.line_spacing = 1.2

# Running header/footer
header = sec.header
p = header.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('吉利试制样车改制数字化项目｜招标技术要求书')
set_run_font(r, size=8.5, color=MUTED)
footer = sec.footer
p = footer.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('第 '); set_run_font(r, size=8.5, color=MUTED)
add_field(p, 'PAGE')
r = p.add_run(' 页'); set_run_font(r, size=8.5, color=MUTED)

def title_page():
    for _ in range(3): doc.add_paragraph('')
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('吉利试制样车改制数字化项目'); set_run_font(r, size=28, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('招标技术要求书'); set_run_font(r, size=24, bold=True, color=DARK)
    p.paragraph_format.space_after = Pt(28)
    for _ in range(3): doc.add_paragraph('')
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(4.0), Cm(8.0)]
    values = [('项目名称','试制样车改制数字化项目'),('文档版本','V15.0（招标技术要求修订稿）')]
    for i,row in enumerate(table.rows):
        for j,c in enumerate(row.cells):
            c.width = widths[j]; c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(c,120,160,120,160)
            p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
            rr=p.add_run(values[i][j]); set_run_font(rr,size=10.5,bold=(j==0),color=BLUE if j==0 else DARK)
    set_table_borders(table, color='D5DEE7')
    doc.add_paragraph('')
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('编制日期：2026年7月'); set_run_font(r,size=10,color=MUTED)
    # A real section break is more robust across Word/Quick Look renderers than
    # an inline page-break run and keeps the cover isolated from the body.
    new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    new_sec.page_width = Cm(21.0)
    new_sec.page_height = Cm(29.7)
    new_sec.top_margin = Cm(2.2)
    new_sec.bottom_margin = Cm(2.0)
    new_sec.left_margin = Cm(2.5)
    new_sec.right_margin = Cm(2.2)
    new_sec.header_distance = Cm(1.0)
    new_sec.footer_distance = Cm(1.0)

PLAIN_REQ_COUNTER = 0

def h1(text):
    global PLAIN_REQ_COUNTER
    PLAIN_REQ_COUNTER = 0
    doc.add_heading(text, level=1)
def h2(text):
    global PLAIN_REQ_COUNTER
    PLAIN_REQ_COUNTER = 0
    doc.add_heading(text, level=2)
def h3(text):
    global PLAIN_REQ_COUNTER
    PLAIN_REQ_COUNTER = 0
    doc.add_heading(text, level=3)
def para(text, bold_lead=None):
    p=doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    if bold_lead and text.startswith(bold_lead):
        r=p.add_run(bold_lead); set_run_font(r,bold=True)
        r=p.add_run(text[len(bold_lead):]); set_run_font(r)
    else:
        r=p.add_run(text); set_run_font(r)
    return p
def bullet(text, level=0):
    p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2')
    p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.2
    r=p.add_run(text); set_run_font(r)
    return p
SHOW_REQ_CODE = True

def req(code, text, priority='P0'):
    global PLAIN_REQ_COUNTER
    p=doc.add_paragraph(style='Requirement')
    if SHOW_REQ_CODE:
        r=p.add_run(f'{code} '); set_run_font(r,bold=True,color=BLUE)
    else:
        PLAIN_REQ_COUNTER += 1
        r=p.add_run(f'{PLAIN_REQ_COUNTER}. '); set_run_font(r,bold=True,color=BLUE)
    r=p.add_run(text); set_run_font(r)
    return p
def table(headers, rows, widths=None, font=9):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    set_table_borders(t)
    for j,h in enumerate(headers):
        c=t.rows[0].cells[j]; set_cell_shading(c,LIGHT); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(c)
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); set_run_font(r,size=font,bold=True,color=BLUE)
    set_repeat_table_header(t.rows[0])
    for vals in rows:
        row=t.add_row(); set_repeat_no_split(row)
        for j,v in enumerate(vals):
            c=row.cells[j]; c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(c)
            p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT if len(str(v))>8 else WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(str(v)); set_run_font(r,size=font)
    if widths:
        for row in t.rows:
            for j,w in enumerate(widths): row.cells[j].width=Cm(w)
    doc.add_paragraph('').paragraph_format.space_after=Pt(1)
    return t

title_page()

h1('1 项目概述')
h2('1.1 项目背景')
para('试制样车改制业务面向研发阶段的功能验证、试验验证、标定和对标等场景，对既有样车实施拆解、更换、加装、切割、焊接、装配、软件调试和质量放行。随着车型研发周期缩短、改制批次和并行项目增多，现有以PPT、Excel、飞书表格、纸质单据和人工协调为主的管理方式，难以持续支撑高频、多车型、多项目和多车间协同。')
para('当前主要问题包括：改制方案和评审意见分散、版本冻结和变更控制不足；计划排产依赖经验，车间、班组、工位和举升机资源状态不透明；物料齐套风险暴露较晚；新装件、拆车件和车辆缺少贯通追溯；质量记录与车辆履历分离；管理层难以及时获取项目、车辆、设备、物料、质量和交付状态。')
h2('1.2 建设目标')
para('本项目应建设以改制任务为驱动、以车辆为核心对象，贯通需求承接、方案评审、计划确认、物料准备、现场执行、质量闭环、客户交付和返修回流的试制样车改制数字化平台，实现以下目标：')
bullet('评审在线：改制方案结构化、附件可标注、意见可闭环、版本可冻结、变更可追溯。')
bullet('物料在线：物料需求、库存齐套、领料配送、车间接收、装车绑定和拆车件去向透明。')
bullet('车辆在线：车辆从需求进入到交付退出全过程状态在线，并形成可回放的一车一档。')
bullet('资源在线：车间、班组、工位和举升机的计划、占用、空闲和超时状态可视。')
bullet('质量在线：接车、拆车、改制、合装、调试、终检、客户验收及返修问题形成闭环。')
bullet('管理在线：项目进度、交付风险、物料风险、设备占用、质量问题和产能指标实时呈现。')
h2('1.3 项目定位')
para('本系统定位为衔接TOCC项目及车辆信息、SAP物料数据、LES仓储物流数据与改制车间执行证据链的专业化业务系统。TOCC作为改制需求、WBS及车辆基础信息的主要来源；SAP提供物料相关数据；LES提供仓库库存、齐套、领料、配送等数据。本系统负责改制业务内部的方案、任务、计划、执行、质量、交付和追溯闭环。')
h2('1.4 建设原则')
for x in ['业务闭环优先：一期优先保障主业务链路完整可运行。','统一模型、柔性配置：统一任务底座支持不同改制类型和复杂度。','少录入、重复用：优先通过系统集成、模板复用和结构化录入减少重复工作。','现场可执行：扫码、确认和报工动作应符合技师实际作业习惯。','人工可干预：系统辅助生成的业务结果应保留人工确认和调整机制。','全程可追溯：关键对象、关键动作、关键版本和关键结论均应留痕。']:
    bullet(x)

h1('2 建设范围与边界')
h2('2.1 业务范围')
table(['业务阶段','本期覆盖内容'],[
('需求承接','TOCC需求接收、WBS关联、车辆范围确认、任务分类及受理'),
('方案评审','结构化方案、附件预览标注、方案信息提取、多人评审、签发冻结及变更控制'),
('计划准备','计划确认、物料齐套、车间/班组/工位/举升机资源确认'),
('改制执行','接车、拆、改、装、调、检等阶段任务和关键节点记录'),
('物料追溯','新件需求与齐套、配送接收、装车绑定、拆车件暂存/回装/报废状态'),
('质量管理','全过程检验、问题单、整改、复验、放行和知识沉淀'),
('交付归档','客户验收、交付签收、一车一档、报告归档和返修回流'),
('经营分析','项目、车辆、产能、设备、物料、质量、交付和异常看板')],[3.4,12.8],9)
h2('2.2 业务对象范围')
para('系统应覆盖SM改制、零星改制、总装改制、车身改制以及零部件/工装开发等业务。系统设计不应简单复制多套割裂流程，应建立统一任务模型，并支持按照任务复杂度选择快速工单、标准改制项目或开发任务等管理模式。')
h2('2.3 组织和现场范围')
para('系统应支持甲方改制业务相关部门、IT管理部门、策划/计划岗位、工艺质量岗位、总装改制车间、车身改制车间、钣金及零部件开发区域、物料管理岗位和项目需求方协同。现场资源应支持车间、工段、班组、工位、举升机、暂存区域、托盘/容器等分层配置。实际组织、车间数量和工位数量以项目调研确认结果为准。')
h2('2.4 系统集成边界')
table(['系统','数据职责','本系统主要动作'],[
('TOCC','改制需求、项目/WBS、车辆基础信息及相关状态','接收需求和基础信息；按确认范围回写任务、车辆或交付状态'),
('SAP','物料主数据及改制物料相关数据','获取物料编码、名称、规格等必要信息，支持改制BOM和物料追溯'),
('LES','库存、齐套、领料、配送及仓库物流状态','查询库存与齐套、发起或接收配送信息、确认车间接收状态')],[2.4,6.4,7.4],9)
SHOW_REQ_CODE = False
h1('3 总体业务要求')
h2('3.1 目标业务主线')
para('目标业务主线为：TOCC需求及WBS进入→任务受理→车辆及范围确认→改制方案结构化→附件标注与方案评审→方案签发冻结→计划确认→物料齐套→车间/班组/工位排产→车辆接收→拆改装调→过程检验与终检→客户验收→交付归档→试验问题及返修回流。')
h2('3.2 方案评审与计划确认分离')
req('BR-001','系统应将“方案评审”和“计划确认”设计为相互关联但相对独立的业务过程。方案评审重点确认完整性、工艺可行性和风险；计划确认重点确认交期、产能、物料和资源可满足性。')
req('BR-002','方案评审通过不应自动等同于交期承诺；计划确认结果应记录甲方要求时间、可承诺时间、数量调整、资源补充及延期协商结论。')
h2('3.3 统一任务模型')
req('BR-003','系统应采用统一任务/订单模型承载不同改制业务，并通过任务类型、流程模板、表单模板、质量模板和计划规则实现差异化管理。')
req('BR-004','快速工单应支持减少非必要审批和操作步骤，但仍须具备需求来源、车辆对象、执行责任、物料变更、质量结论和完成记录。')
req('BR-005','标准改制项目应支持一个项目关联多台车辆，并支持每台车辆独立排产、执行、质量和交付。')
req('BR-006','零部件或工装开发任务应支持独立发起，也应支持关联一个或多个改制项目。')
h2('3.4 改制阶段模型')
para('系统默认应支持接车、拆、改、装、调、检验和交付等一级阶段。考虑改制业务柔性较高，投标人不得强制将所有现场动作拆分为固定流水工序；应支持按任务类型配置阶段、子任务、检查项和关键扫码节点。')
h2('3.5 标准周期与异常')
para('系统应支持按业务类型配置阶段标准周期。当前可将“拆约3天、装约4天、调试/换装约3天、整体约10天”作为初始参考值，最终以甲方确认规则为准。超过计划节点或标准周期时，系统应预警并要求记录原因。')

SHOW_REQ_CODE = True
h1('4 功能需求')
h2('4.1 生产基础数据与资源建模')
for code,text,pri in [
('FR-001','系统应支持工厂、车间、工段、班组、工位和工作中心等生产组织及层级关系建模。','P0'),
('FR-002','系统应支持举升机、区域、库位、托盘及容器等生产资源和现场对象的基础数据管理。','P0'),
('FR-003','系统应支持车型、任务类型、改制阶段、物料追溯颗粒度和质量检查模板等业务基础数据配置。','P0')]: req(code,text,pri)

h2('4.2 TOCC需求、WBS与任务管理')
for code,text,pri in [
('FR-101','系统应从TOCC接收改制需求、WBS及车辆基础信息，避免重复录入。','P0'),
('FR-102','系统应保留TOCC来源标识、原始单号、同步时间、数据版本和同步结果。','P0'),
('FR-103','系统应支持将一个WBS关联一个或多个改制任务，并按WBS统计车辆、任务、物料和人天等基础数据。','P0'),
('FR-104','系统应支持需求受理、退回补充、任务分类、责任人分配和优先级设置。','P0'),
('FR-105','系统应支持快速工单、标准改制项目和开发任务等模式，并支持甲方配置分类规则。','P0'),
('FR-106','系统应支持一个项目关联多台车辆，每台车辆形成独立执行和交付状态。','P0'),
('FR-107','系统应支持任务暂停、取消、恢复、变更和关闭，并记录原因及审批记录。','P0'),
('FR-108','系统应支持任务台账、组合查询、批量导出和个人待办。','P0')]: req(code,text,pri)

h2('4.3 车辆建档与一车一档')
for code,text,pri in [
('FR-201','系统应通过TOCC车辆ID、VIN、样车编号等建立车辆唯一身份，避免重复建档。','P0'),
('FR-202','系统应记录车型、配置、来源、所属项目、当前状态、所在位置和责任任务等信息。','P0'),
('FR-203','系统应支持同一车辆多次改制，每次改制形成独立履历并按时间轴汇总。','P0'),
('FR-204','系统应展示车辆关联的方案版本、阶段记录、工位设备、物料变更、质量问题、图片附件和交付文件。','P0'),
('FR-205','系统应支持车辆准入校验及异常提示，包括身份缺失、状态冲突、已有未完结任务等。','P1'),
('FR-206','系统应支持车辆状态查询、二维码生成和打印。','P0'),
('FR-207','系统应支持按车辆生成一车一档报告，并支持在线查看和导出。','P0')]: req(code,text,pri)

h2('4.4 结构化方案与在线评审')
for code,text,pri in [
('FR-301','系统应支持按改制类型配置结构化方案模板，字段至少覆盖改制目的、车辆范围、改前状态、改后目标、改制部位、拆除件、新装件、回装件、工装、软件调试、质量要求、风险和计划节点。','P0'),
('FR-302','系统应支持结构化字段与原始附件共同构成完整改制方案，附件不得脱离任务和版本单独流转。','P0'),
('FR-303','系统应支持PPT/PPTX、PDF、Word、Excel及图片等文件上传、在线预览和下载。','P0'),
('FR-304','系统应支持在PPT、PDF和图片页面上进行框选、箭头、画笔、高亮、文字意见和定位标注。','P0'),
('FR-305','标注意见应支持责任人、回复、处理状态、关闭结论和评审轮次，并与附件页码、坐标位置及版本关联。','P0'),
('FR-306','系统应支持多人会签、顺序评审、并行评审、退回修改、再次提交、批准和签发。','P0'),
('FR-307','系统应支持版本自动编号、版本说明、版本差异、评审记录和历史版本查看。','P0'),
('FR-308','方案签发后应形成冻结执行版本；冻结版本不得被直接覆盖，变更须发起受控流程。','P0'),
('FR-309','方案变更应提示对车辆范围、物料需求、排产计划、已开工任务和质量要求的潜在影响。','P1'),
('FR-310','生产、物料和质量记录应引用明确的方案版本，确保执行依据可追溯。','P0')]: req(code,text,pri)

h2('4.5 计划与排产管理')
for code,text,pri in [
('FR-401','系统应支持记录需求方要求交期、内部评估交期、承诺数量、协商结论和责任人。','P0'),
('FR-402','系统应支持年、月、周、日多层级计划，并支持项目、车辆和阶段三个层级的计划展示。','P0'),
('FR-403','系统应将车间、班组、工位、举升机、标准工期、物料齐套和前后阶段关系作为排产约束。','P0'),
('FR-404','系统应支持自动生成可执行的推荐排程，并说明关键冲突或风险。','P0'),
('FR-405','系统应支持锁定任务、插单、拖拽调整、换班组、换工位、重新计算和冲突校验。','P0'),
('FR-406','系统应保留人工调整入口，记录调整前后计划、调整人、时间和原因。','P0'),
('FR-407','系统应支持设置车间间不可随意跨排、车间内工段或工位可柔性调整等约束规则。','P0'),
('FR-408','系统应支持按班组或班组长派工，不强制将所有任务预排至单个技师。','P0'),
('FR-409','系统应识别资源冲突、物料不齐、计划超期、设备超时和任务依赖异常并预警。','P0')]: req(code,text,pri)

h2('4.6 生产执行管理')
for code,text,pri in [
('FR-501','系统应支持将已确认的排产计划转换为车辆级和阶段级生产任务，并下达到相应车间及班组。','P0'),
('FR-502','系统应支持按照车间、班组、工位及举升机进行生产任务派工和资源分配。','P0'),
('FR-503','系统应支持生产任务的开工、暂停、恢复、完工和转序，并记录各环节实际开始及完成时间。','P0'),
('FR-504','系统应支持车辆开工时与工位、举升机建立占用关系，完工或转序后释放相应资源。','P0'),
('FR-505','系统应实时展示车辆当前生产阶段、所在车间、责任班组、工位、执行状态及在制时长。','P0'),
('FR-506','系统应向现场人员展示当前生产任务、车辆信息、冻结方案版本、作业要求、所需物料和质量检查内容。','P0'),
('FR-507','系统应支持通过扫码或人工确认方式记录车辆上线、阶段开始、阶段完成、转序和下线等关键生产节点。','P0'),
('FR-508','系统应支持上传生产过程中的图片、文件及相关记录，并与车辆、生产任务和生产阶段关联。','P1'),
('FR-509','系统应支持缺料、方案变更、质量问题、任务暂停、工位占用超时等生产异常的反馈和跟踪。','P0'),
('FR-510','系统应支持返工返修任务重新下达、重新派工和执行过程跟踪。','P0')]: req(code,text,pri)

h2('4.7 生产物料与齐套管理')
for code,text,pri in [
('FR-601','系统应基于冻结方案或改制BOM形成任务和车辆物料需求。','P0'),
('FR-602','系统应从SAP获取必要的物料主数据，并保留来源及同步记录。','P0'),
('FR-603','系统应与LES交互库存、齐套、领料、配送和仓库状态。','P0'),
('FR-604','系统应按任务、车辆或批次展示物料齐套率、缺料清单、预计到料时间和风险状态。','P0'),
('FR-605','系统应支持车间接收核对，记录应收、实收、差异、接收人和接收时间。','P0'),
('FR-606','系统应支持关键新件扫码装车，记录物料编码、批次/序列号、装车车辆、时间和操作人。','P0'),
('FR-607','系统应支持按单件、批次、箱或托盘等不同颗粒度配置追溯方式。','P0'),
('FR-608','一期可按任务精准领料并在任务结束后确认消耗；系统应预留余料、退料和差异处理能力。','P1'),
('FR-609','物料不齐时，系统应对排产和开工进行提示或按规则控制。','P0')]: req(code,text,pri)

h2('4.8 拆换件与暂存管理')
for code,text,pri in [
('FR-701','系统应记录拆车件来源车辆、改制任务、拆卸阶段、拆卸时间及责任班组。','P0'),
('FR-702','系统应支持按单件、批次、箱、托盘或容器管理拆车件，颗粒度可配置。','P0'),
('FR-703','系统应支持托盘/容器编码、打印、绑定、拆分、合并和状态管理。','P0'),
('FR-704','系统应记录托盘/容器所在区域、库位及移动记录。','P0'),
('FR-705','拆车件状态应至少包括待处理、暂存、待回装、已回装、报废和其他处置。','P0'),
('FR-706','回装时应支持核验拆车件与来源车辆、任务和方案版本的关联关系。','P0'),
('FR-707','报废业务可由外部业务流程处理，本系统应记录业务处置状态和必要凭证。','P1'),
('FR-708','投标人应通过现场调研确定关键件单件管理与普通件容器管理的划分规则。','P0')]: req(code,text,pri)

h2('4.9 生产质量管理')
for code,text,pri in [
('FR-801','系统应支持接车检、拆车过程/终检、钣金接车检、钣金过程/终检、合装过程/终检、调试结论、下线终检和客户验收等质量节点。','P0'),
('FR-802','质量节点和检查模板应可按任务类型、车型和阶段配置。','P0'),
('FR-803','系统应支持检查项、结果、测量值、图片附件、结论和签字确认。','P0'),
('FR-804','系统应支持质量问题提出、分类、责任分配、原因分析、整改措施、复验和关闭。','P0'),
('FR-805','质量问题应关联项目、WBS、任务、车辆、阶段、方案版本、物料及责任组织。','P0'),
('FR-806','系统应支持不合格阻断、让步放行和授权审批，具体质量门禁规则由甲方配置。','P0'),
('FR-807','系统应支持返工返修任务生成及重新排产。','P0'),
('FR-808','系统应支持质量问题统计、重复问题识别、横向展开和知识沉淀。','P1')]: req(code,text,pri)

h2('4.10 完工与交付管理')
for code,text,pri in [
('FR-901','系统应支持车辆完成内部检验后进入待交付状态，并记录客户验收和提车信息。','P0'),
('FR-902','系统应支持客户验收结论、遗留问题、签收人、签收时间及附件。','P0'),
('FR-903','系统应支持交付后试验问题反馈，并根据问题创建返修任务或重新进入改制流程。','P0'),
('FR-904','系统应自动汇总需求、方案、计划、执行、物料、质量和交付数据形成交付包。','P0'),
('FR-905','系统应按照接口确认范围向TOCC回写任务、车辆或交付状态。','P0'),
('FR-906','系统应支持项目、车辆和时间维度的历史归档、查询和导出。','P0')]: req(code,text,pri)

h2('4.11 生产监控与分析')
for code,text,pri in [
('FR-1001','系统应提供项目总览，展示任务、车辆、计划、完成率、风险和交付状态。','P0'),
('FR-1002','系统应提供车间看板，展示各工位/举升机的空闲、占用、车辆、阶段和超时状态。','P0'),
('FR-1003','系统应提供物料齐套和缺料风险看板。','P0'),
('FR-1004','系统应提供质量问题、整改进度、逾期和重复问题分析。','P0'),
('FR-1005','系统应支持计划工期与实际工期、标准周期达成率、设备利用率和车辆在制周期分析。','P0'),
('FR-1006','系统应支持按WBS、项目、车型、车间、班组、任务类型和时间进行统计。','P0'),
('FR-1007','系统应支持超期、缺料、长时间占用、质量未闭环、方案变更和接口异常等消息预警。','P0'),
('FR-1008','看板指标、口径、阈值和展示范围应可配置。','P1')]: req(code,text,pri)

SHOW_REQ_CODE = False
h1('5 系统集成要求')
h2('5.1 总体要求')
para('系统应围绕改制业务与TOCC、SAP和LES进行数据协同，减少重复录入，保证需求、车辆、物料、库存、齐套、配送和业务状态能够在相关环节使用。具体数据范围以业务流程确认结果为准。')
h2('5.2 TOCC接口')
for code,text in [
('IF-101','接收改制需求、WBS、项目和车辆基础信息。'),
('IF-102','支持需求变更、撤销及车辆范围变化的同步识别。'),
('IF-103','按双方确认范围回写任务受理、执行、完成、交付或履历信息。'),
('IF-104','支持来源数据版本、同步时间、处理状态和失败原因查询。')]: req(code,text,'P0')
h2('5.3 SAP接口')
for code,text in [('IF-201','获取物料编码、名称、规格、单位、状态等必要主数据。'),('IF-202','支持改制BOM或物料需求所需数据交互，具体数据主责以接口设计为准。'),('IF-203','物料主数据变更后应支持增量同步及异常提示。')]: req(code,text,'P0')
h2('5.4 LES接口')
for code,text in [('IF-301','查询库存、可用量、齐套状态及必要的预计到料信息。'),('IF-302','支持领料/配送需求及仓库处理状态交互。'),('IF-303','接收配送明细并支持车间实物接收确认。'),('IF-304','支持差异、取消、退回或异常状态的处理和对账。')]: req(code,text,'P0')

h1('6 技术要求')
h2('6.1 一体化平台基础能力')
for code,text in [
('TP-001','系统应提供统一应用门户，对改制业务相关功能模块进行集中访问和统一导航。'),
('TP-002','系统应提供个人工作台，集中展示个人待办、已办、业务提醒和常用功能快捷入口，并支持用户配置常用入口。'),
('TP-003','系统应支持用户、部门、岗位、业务角色和用户分组等平台基础信息的统一维护，并供业务流程和任务分配使用。'),
('TP-004','系统应支持应用模块、菜单层级和业务功能入口的配置，以适应不同业务角色的工作入口。'),
('TP-005','系统应支持公共数据字典的统一维护，为业务状态、任务类型、问题类型和其他标准选项提供一致的数据来源。'),
('TP-006','系统应支持业务对象和单据编码规则配置，可按前缀、日期、流水号及业务属性组合生成编码。'),
('TP-007','系统应提供动态表单配置能力，支持配置字段、布局、校验规则、默认值及表单与业务数据的关联关系。'),
('TP-008','系统应提供可视化流程配置能力，支持配置流程节点、流转路径、办理人、会签方式、条件分支及审批规则。'),
('TP-009','系统应提供统一流程任务中心，支持流程发起、我的申请、待办、已办、抄送、撤回、转办、委派及批注等业务操作。'),
('TP-010','系统应提供消息中心，支持站内消息、业务提醒和通知公告，并记录消息发送及阅读状态。'),
('TP-011','系统应提供统一文件中心，支持业务附件上传、分类、预览、下载和与业务对象关联。'),
('TP-012','系统应提供可视化页面和看板配置能力，支持通过图表、表格、文本等组件配置业务看板，并绑定业务数据。')]: req(code,text,'P0')

h2('6.2 一体化集成平台能力')
for code,text in [
('IP-001','系统应提供统一的应用集成管理能力，集中维护已接入应用、系统标识、连接信息和接口目录。'),
('IP-002','系统应提供连接器配置能力，支持针对不同外部系统配置连接方式、调用地址、参数和数据对应关系。'),
('IP-003','系统应支持通过RESTful API、WebService等通用方式与外部应用进行数据交互，并可根据业务需要扩展其他集成方式。'),
('IP-004','系统应提供接口在线配置和开发能力，支持接口新增、修改、停用、发布及相关处理逻辑维护。'),
('IP-005','系统应支持接口字段映射、数据格式转换、数据合并拆分和基础数据处理规则配置。'),
('IP-006','系统应提供接口服务目录及接口文档管理，支持查看接口地址、输入输出参数、调用说明和数据示例。'),
('IP-007','系统应支持对TOCC、SAP、LES等外部系统形成可复用的连接配置，便于后续新增接口和扩展集成范围。'),
('IP-008','系统应支持查看接口调用记录、处理状态和异常信息，并对失败数据进行重新处理。')]: req(code,text,'P0')

h2('6.3 开发技术栈要求')
for code,text in [
('TS-001','系统应采用B/S架构及前后端分离的技术模式，前后端独立开发、部署，通过标准网络协议通信，并兼容主流浏览器，保障访问便捷与可维护性。'),
('TS-002','系统前端应采用Vue 3技术体系，基于Vite 3构建，使用Pinia进行状态管理，结合Element Plus组件库与Axios请求库等技术栈开发。'),
('TS-003','系统后端应采用Java技术体系，基于Spring Boot 3.5、Spring Security等成熟、稳定的技术组件进行构建，确保高安全性与可扩展性。'),
('TS-004','前后端业务交互应采用标准化接口方式，以RESTful API为规范，接口数据统一使用JSON格式，确保交互通用、简洁、易集成。'),
('TS-005','系统功能采用模块化设计，平台基础能力、业务功能与外部系统集成应保持合理分层。系统建设统一的用户中心集中管理用户身份，依托一体化门户集成多个业务应用，基于用户中心与权限中心实现应用免登与统一授权。平台应提供统一的任务调度（Job）、流程、消息、文件、日志等公共中心能力，为各应用运行提供一致支撑，便于业务灵活扩展与能力复用。')]: req(code,text,'P0')

# Keep headings with following content and set fonts for any auto-created list text.
for p in doc.paragraphs:
    if p.style.name.startswith('Heading'):
        p.paragraph_format.keep_with_next = True
    for r in p.runs:
        if not r.font.name:
            set_run_font(r)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
